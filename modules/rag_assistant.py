"""
RAG-based Learning Assistant Module - Milestone 5
Integrates document upload and question-answering using FAISS + HuggingFace
Allows users to upload PDFs/notes and ask questions from documents
"""

import os
import uuid
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime
import numpy as np

# Document processing
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
import pypdf
from pypdf import PdfReader

try:
    import docx
except ImportError:
    docx = None

# Vector database and embeddings
try:
    import faiss
except ImportError:
    faiss = None

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

# Text processing
from dataclasses import dataclass
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import Config
from .qa_system import QASystem

@dataclass
class Document:
    page_content: str
    metadata: dict

logger = logging.getLogger(__name__)

class RAGLearningAssistant:
    """RAG-based Learning Assistant for document-based Q&A"""
    
    def __init__(self):
        self.qa_system = QASystem()
        self.embedding_model = None
        self.vector_store = None
        self.document_store = {}
        self.document_chunks = {}
        self.next_vector_index = 0
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", " ", ""]
        )
        self._initialize_embedding_model()
    
    def _initialize_embedding_model(self):
        """Initialize the sentence transformer model for embeddings"""
        if SentenceTransformer is None:
            logger.error("sentence-transformers package is not installed. Embedding model unavailable.")
            self.embedding_model = None
            self.vector_store = None
            return

        if faiss is None:
            logger.error("faiss package is not installed. Vector store unavailable.")
            self.embedding_model = None
            self.vector_store = None
            return

        try:
            # Use a lightweight but effective model
            model_name = 'all-MiniLM-L6-v2'
            self.embedding_model = SentenceTransformer(model_name)
            logger.info(f"Initialized embedding model: {model_name}")
            
            # Initialize FAISS index
            embedding_dim = self.embedding_model.get_sentence_embedding_dimension()
            self.vector_store = faiss.IndexFlatIP(embedding_dim)  # Inner product for similarity
            
        except Exception as e:
            logger.error(f"Failed to initialize embedding model: {str(e)}")
            self.embedding_model = None
            self.vector_store = None
    
    def upload_document(self, file_path: str, file_name: str = None) -> Dict[str, Any]:
        """
        Upload and process a document for RAG
        
        Args:
            file_path (str): Path to the document file
            file_name (str): Optional custom name for the document
            
        Returns:
            dict: Upload result with document ID and metadata
        """
        try:
            if not os.path.exists(file_path):
                return {
                    'success': False,
                    'error': 'File not found',
                    'document_id': None
                }
            
            # Check file extension
            file_extension = os.path.splitext(file_path)[1].lower()
            if file_extension not in ['.pdf', '.txt', '.docx']:
                return {
                    'success': False,
                    'error': 'Unsupported file format. Use PDF, TXT, or DOCX',
                    'document_id': None
                }
            
            # Extract text from document
            text_content = self._extract_text_from_file(file_path, file_extension)
            if not text_content.strip():
                return {
                    'success': False,
                    'error': 'No text content found in document',
                    'document_id': None
                }
            
            # Generate document ID
            document_id = str(uuid.uuid4())
            display_name = file_name or os.path.basename(file_path)
            
            # Split text into chunks
            documents = self._split_text_into_chunks(text_content, document_id, display_name)
            
            # Generate embeddings and store
            embedding_result = self._store_document_embeddings(documents, document_id)
            if not embedding_result['success']:
                return embedding_result
            
            # Store document metadata
            self.document_store[document_id] = {
                'id': document_id,
                'name': display_name,
                'file_path': file_path,
                'file_type': file_extension,
                'upload_date': datetime.now().isoformat(),
                'content_length': len(text_content),
                'num_chunks': len(documents),
                'processed': True
            }
            
            return {
                'success': True,
                'document_id': document_id,
                'document_name': display_name,
                'chunks_created': len(documents),
                'content_length': len(text_content),
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Error uploading document: {str(e)}")
            return {
                'success': False,
                'error': f'Error processing document: {str(e)}',
                'document_id': None
            }
    
    def _extract_text_from_file(self, file_path: str, file_extension: str) -> str:
        """Extract text content from different file types"""
        try:
            if file_extension == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_extension == '.txt':
                return self._extract_text_from_txt(file_path)
            elif file_extension == '.docx':
                return self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file extension: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            if fitz is not None:
                try:
                    doc = fitz.open(file_path)
                    for page in doc:
                        text += page.get_text()
                    doc.close()
                except Exception as e:
                    logger.warning(f"PyMuPDF failed on {file_path}: {str(e)}")
                    text = ""
            
            if text.strip():
                return text
            
            # Fallback to pypdf if PyMuPDF is unavailable or fails
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
                return text
                
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as file:
                return file.read().decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if docx is None:
            logger.error("python-docx is not installed, DOCX file reading is unavailable")
            return ""

        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            return ""
    
    def _split_text_into_chunks(self, text: str, document_id: str, document_name: str) -> List[Document]:
        """Split text into chunks for embedding"""
        try:
            # Create langchain documents
            doc = Document(
                page_content=text,
                metadata={
                    'document_id': document_id,
                    'document_name': document_name,
                    'source': document_name
                }
            )
            
            # Split into chunks
            chunks = self.text_splitter.split_documents([doc])
            
            # Add chunk-specific metadata
            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    'chunk_id': f"{document_id}_chunk_{i}",
                    'chunk_index': i,
                    'total_chunks': len(chunks)
                })
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error splitting text: {str(e)}")
            return []
    
    def _store_document_embeddings(self, documents: List[Document], document_id: str) -> Dict[str, Any]:
        """Generate embeddings and store in FAISS vector database"""
        try:
            if not self.embedding_model or not self.vector_store:
                return {
                    'success': False,
                    'error': 'Embedding model not available'
                }
            
            texts = [doc.page_content for doc in documents]
            if not texts:
                return {
                    'success': False,
                    'error': 'No text content to embed'
                }
            
            embeddings = self.embedding_model.encode(texts)
            embeddings = embeddings / np.linalg.norm(embeddings, axis=1, keepdims=True)
            self.vector_store.add(embeddings)
            
            start_idx = self.next_vector_index
            for i, doc in enumerate(documents):
                idx = start_idx + i
                self.document_chunks[idx] = {
                    'content': doc.page_content,
                    'metadata': doc.metadata,
                    'document_id': document_id,
                    'embedding_index': idx
                }
            self.next_vector_index += len(documents)
            
            return {
                'success': True,
                'embeddings_created': len(embeddings),
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Error storing embeddings: {str(e)}")
            return {
                'success': False,
                'error': f'Error storing embeddings: {str(e)}'
            }
    
    def ask_document_question(self, question: str, document_ids: Optional[List[str]] = None, 
                            top_k: int = 5) -> Dict[str, Any]:
        """
        Ask a question about uploaded documents using RAG
        
        Args:
            question (str): Question to ask
            document_ids (List[str]): Specific documents to search (None for all)
            top_k (int): Number of relevant chunks to retrieve
            
        Returns:
            dict: Answer with sources and confidence
        """
        try:
            if not question.strip():
                return {
                    'success': False,
                    'error': 'Question cannot be empty',
                    'answer': None,
                    'sources': []
                }
            
            if not self.embedding_model or not self.vector_store:
                return {
                    'success': False,
                    'error': 'RAG system not initialized',
                    'answer': None,
                    'sources': []
                }
            
            if not hasattr(self, 'document_chunks') or not self.document_chunks:
                return {
                    'success': False,
                    'error': 'No documents uploaded yet',
                    'answer': None,
                    'sources': []
                }
            
            # Retrieve relevant chunks
            relevant_chunks = self._retrieve_relevant_chunks(question, document_ids, top_k)
            
            if not relevant_chunks:
                return {
                    'success': False,
                    'error': 'No relevant information found in documents',
                    'answer': None,
                    'sources': []
                }
            
            # Build context from relevant chunks
            context = self._build_context_from_chunks(relevant_chunks)
            
            # Generate answer using QA system with context
            response = self.qa_system.ask_question(
                question=question,
                context=f"Based on the following document excerpts:\n\n{context}",
                provider="auto"
            )
            
            if not response['success']:
                return {
                    'success': False,
                    'error': response['error'],
                    'answer': None,
                    'sources': []
                }
            
            # Prepare sources information
            sources = []
            for chunk in relevant_chunks:
                sources.append({
                    'document_name': chunk['metadata'].get('document_name', 'Unknown'),
                    'chunk_index': chunk['metadata'].get('chunk_index', 0),
                    'content_preview': chunk['content'][:200] + "..." if len(chunk['content']) > 200 else chunk['content'],
                    'relevance_score': chunk.get('score', 0.0)
                })
            
            return {
                'success': True,
                'answer': response['answer'],
                'sources': sources,
                'context_used': context,
                'provider_used': response['provider_used'],
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Error processing document question: {str(e)}")
            return {
                'success': False,
                'error': f'Error processing question: {str(e)}',
                'answer': None,
                'sources': []
            }
    
    def _retrieve_relevant_chunks(self, question: str, document_ids: Optional[List[str]], 
                                top_k: int) -> List[Dict]:
        """Retrieve most relevant document chunks for the question"""
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode([question])
            query_embedding = query_embedding / np.linalg.norm(query_embedding, axis=1, keepdims=True)
            
            # Search in vector store
            scores, indices = self.vector_store.search(query_embedding, top_k * 2)  # Get more candidates
            
            # Filter and rank results
            relevant_chunks = []
            for i, (score, idx) in enumerate(zip(scores[0], indices[0])):
                if idx == -1:  # FAISS returns -1 for empty slots
                    continue
                
                if idx in self.document_chunks:
                    chunk_data = self.document_chunks[idx].copy()
                    
                    # Filter by document IDs if specified
                    if document_ids and chunk_data['document_id'] not in document_ids:
                        continue
                    
                    chunk_data['score'] = float(score)
                    relevant_chunks.append(chunk_data)
                    
                    if len(relevant_chunks) >= top_k:
                        break
            
            # Sort by relevance score (descending)
            relevant_chunks.sort(key=lambda x: x['score'], reverse=True)
            
            return relevant_chunks[:top_k]
            
        except Exception as e:
            logger.error(f"Error retrieving relevant chunks: {str(e)}")
            return []
    
    def _build_context_from_chunks(self, chunks: List[Dict]) -> str:
        """Build context string from retrieved chunks"""
        context_parts = []
        
        for i, chunk in enumerate(chunks, 1):
            document_name = chunk['metadata'].get('document_name', 'Unknown Document')
            content = chunk['content'].strip()
            
            context_parts.append(f"[Source {i} - {document_name}]\n{content}\n")
        
        return "\n".join(context_parts)
    
    def get_uploaded_documents(self) -> List[Dict[str, Any]]:
        """Get list of all uploaded documents"""
        return [doc_info for doc_info in self.document_store.values()]
    
    def delete_document(self, document_id: str) -> Dict[str, Any]:
        """Delete a document and its embeddings"""
        try:
            if document_id not in self.document_store:
                return {
                    'success': False,
                    'error': 'Document not found'
                }
            
            doc_info = self.document_store.pop(document_id)
            if hasattr(self, 'document_chunks'):
                chunks_to_remove = [
                    idx for idx, chunk in self.document_chunks.items()
                    if chunk['document_id'] == document_id
                ]
                for idx in chunks_to_remove:
                    del self.document_chunks[idx]
            
            return {
                'success': True,
                'message': f'Document "{doc_info["name"]}" deleted successfully',
                'document_name': doc_info['name']
            }
            
        except Exception as e:
            logger.error(f"Error deleting document: {str(e)}")
            return {
                'success': False,
                'error': f'Error deleting document: {str(e)}'
            }
    
    def get_system_status(self) -> Dict[str, Any]:
        """Get RAG system status and statistics"""
        return {
            'embedding_model_loaded': self.embedding_model is not None,
            'vector_store_initialized': self.vector_store is not None,
            'total_documents': len(self.document_store),
            'total_chunks': len(getattr(self, 'document_chunks', {})),
            'vector_store_size': self.vector_store.ntotal if self.vector_store else 0,
            'supported_formats': ['.pdf', '.txt', '.docx'],
            'max_file_size_mb': Config.MAX_CONTENT_LENGTH / (1024 * 1024)
        }